# smartjudi2/scripts/load_legal_data.py

import os
import requests
import logging
from dotenv import load_dotenv
from typing import List, Dict, Any

# For PDF and DOCX parsing
try:
    from pypdf import PdfReader  # Use pypdf instead of PyPDF2 for modern Python
except ImportError:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        print("Warning: pypdf or PyPDF2 not installed. PDF parsing will not work.")
        PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    print("Warning: python-docx not installed. DOCX parsing will not work.")
    DocxDocument = None

# For text splitting
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
except ImportError:
    print("Warning: langchain not installed. Text splitting will use simple method.")
    RecursiveCharacterTextSplitter = None

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Load environment variables from .env file
load_dotenv()
RAG_API_URL = os.getenv("RAG_API_URL")

if not RAG_API_URL:
    logger.error("RAG_API_URL environment variable not set. Exiting.")
    exit(1)

# Initialize HTTP client for RAG API
rag_api_client = requests.Session()


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extracts text from a PDF file.
    يستخرج النص من ملف PDF.
    """
    if PdfReader is None:
        logger.error("PDF reader not available. Install pypdf or PyPDF2.")
        return ""
    
    text = ""
    try:
        reader = PdfReader(pdf_path)
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF {pdf_path}: {e}")
        return ""


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extracts text from a DOCX file.
    يستخرج النص من ملف DOCX.
    """
    if DocxDocument is None:
        logger.error("DOCX reader not available. Install python-docx.")
        return ""
    
    text = ""
    try:
        doc = DocxDocument(docx_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {docx_path}: {e}")
        return ""


def load_documents_from_directory(directory_path: str) -> List[Dict[str, Any]]:
    """
    Loads and extracts text from PDF, DOCX, and TXT files in a given directory.
    يقوم بتحميل واستخراج النصوص من ملفات PDF و DOCX و TXT في دليل معين.
    """
    documents = []
    for root, _, files in os.walk(directory_path):
        for file_name in files:
            file_path = os.path.join(root, file_name)
            content = ""
            if file_name.endswith(".pdf"):
                content = extract_text_from_pdf(file_path)
            elif file_name.endswith(".docx"):
                content = extract_text_from_docx(file_path)
            elif file_name.endswith(".txt"):
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                except Exception as e:
                    logger.error(f"Error reading TXT file {file_path}: {e}")
                    content = ""
            
            if content:
                documents.append({
                    "page_content": content,
                    "metadata": {"source": file_path, "file_name": file_name}
                })
                logger.info(f"Loaded document: {file_name}")
    return documents


def chunk_documents(documents: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Chunks documents into smaller pieces with overlap using RecursiveCharacterTextSplitter.
    يقسم المستندات إلى أجزاء أصغر مع تداخل باستخدام RecursiveCharacterTextSplitter.
    """
    if RecursiveCharacterTextSplitter:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            add_start_index=True,
        )
        chunked_docs = []
        for doc in documents:
            try:
                from langchain_core.documents import Document as LangchainDocument
                langchain_doc = LangchainDocument(
                    page_content=doc["page_content"],
                    metadata=doc["metadata"]
                )
                chunks = text_splitter.split_documents([langchain_doc])
                for chunk in chunks:
                    chunked_docs.append({
                        "page_content": chunk.page_content,
                        "metadata": chunk.metadata
                    })
            except Exception as e:
                logger.error(f"Error chunking document: {e}")
                # Fallback: add document as-is if chunking fails
                chunked_docs.append(doc)
        logger.info(f"Chunked {len(documents)} documents into {len(chunked_docs)} chunks.")
        return chunked_docs
    else:
        # Simple chunking without langchain
        logger.warning("Using simple chunking method (langchain not available).")
        chunked_docs = []
        chunk_size = 1000
        chunk_overlap = 200
        for doc in documents:
            content = doc["page_content"]
            metadata = doc["metadata"]
            start = 0
            while start < len(content):
                end = start + chunk_size
                chunk = content[start:end]
                chunked_docs.append({
                    "page_content": chunk,
                    "metadata": {**metadata, "chunk_start": start}
                })
                start = end - chunk_overlap
        logger.info(f"Chunked {len(documents)} documents into {len(chunked_docs)} chunks (simple method).")
        return chunked_docs


def index_documents_to_rag(chunked_documents: List[Dict[str, Any]]) -> bool:
    """
    Indexes chunked documents into the Hugging Face RAG API.
    يفهرس المستندات المقسمة إلى واجهة برمجة تطبيقات RAG المستضافة على Hugging Face.
    """
    if not RAG_API_URL:
        logger.error("RAG_API_URL is not set. Cannot index documents.")
        return False

    add_docs_url = f"{RAG_API_URL}/add_documents"
    
    # Process in batches to avoid overwhelming the API
    batch_size = 50
    total_batches = (len(chunked_documents) + batch_size - 1) // batch_size
    
    for i in range(0, len(chunked_documents), batch_size):
        batch = chunked_documents[i:i + batch_size]
        batch_num = (i // batch_size) + 1
        logger.info(f"Indexing batch {batch_num}/{total_batches} ({len(batch)} documents)...")
        
        try:
            response = rag_api_client.post(add_docs_url, json=batch, timeout=120)
            response.raise_for_status()
            logger.info(f"Successfully indexed batch {batch_num}: {response.json()}")
        except requests.exceptions.RequestException as e:
            logger.error(f"Failed to index batch {batch_num} to RAG engine: {e}")
            if hasattr(e, 'response') and e.response is not None:
                logger.error(f"RAG API Error Response: {e.response.text}")
            return False
    
    logger.info(f"Successfully indexed all {len(chunked_documents)} chunks to RAG engine.")
    return True


def main(data_directory: str):
    """
    Main function to load, chunk, and index legal data.
    الدالة الرئيسية لتحميل وتقسيم وفهرسة البيانات القانونية.
    """
    logger.info(f"Starting legal data loading process from: {data_directory}")
    
    if not os.path.exists(data_directory):
        logger.error(f"Directory {data_directory} does not exist.")
        return
    
    # 1. Load documents
    # 1. تحميل المستندات
    raw_documents = load_documents_from_directory(data_directory)
    if not raw_documents:
        logger.warning("No documents found to process. Exiting.")
        return

    # 2. Chunk documents
    # 2. تقسيم المستندات
    chunked_documents = chunk_documents(raw_documents)

    # 3. Index documents to RAG engine
    # 3. فهرسة المستندات إلى محرك RAG
    if index_documents_to_rag(chunked_documents):
        logger.info("Legal data indexing completed successfully.")
    else:
        logger.error("Legal data indexing failed.")


if __name__ == "__main__":
    # Example usage: specify the directory containing your legal documents
    # مثال على الاستخدام: حدد الدليل الذي يحتوي على مستنداتك القانونية
    # Make sure to create this directory and place your PDF, DOCX, TXT files inside.
    # تأكد من إنشاء هذا الدليل ووضع ملفات PDF و DOCX و TXT بداخله.
    LEGAL_DATA_DIR = os.getenv("LEGAL_DATA_DIR", "./legal_documents")
    if not os.path.exists(LEGAL_DATA_DIR):
        os.makedirs(LEGAL_DATA_DIR)
        logger.info(f"Created directory: {LEGAL_DATA_DIR}. Please place your legal documents here.")
    
    main(LEGAL_DATA_DIR)
